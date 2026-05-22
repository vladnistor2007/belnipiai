import os
import io
import re
import difflib
import base64
import json
import time
import hmac
import hashlib
import tempfile
import uuid
import threading
import urllib.request
from concurrent.futures import ThreadPoolExecutor, as_completed
from flask import Flask, request, Response, send_from_directory, stream_with_context, session, redirect, jsonify
import ollama
from rag_engine import (
    get_db,
    index_document as rag_index_document,
    invalidate_kb_caches,
    hybrid_search,
    expand_with_neighbors,
    build_context_chunks,
    fetch_all_context_pages,
    build_rag_context,
)
from doc_generator import (
    SYSTEM_PROMPT_DOCGEN,
    is_document_request,
    detect_format,
    make_output_path,
    extract_code,
    run_code,
    TEMP_DIR as DOC_TEMP_DIR,
)
from rag_chunker import chunk_text as rag_chunk_text, chunk_pages as rag_chunk_pages, chunk_excel as rag_chunk_excel

app = Flask(__name__, static_folder="static")
ROOT_DIR = app.root_path
app.secret_key = os.environ.get("AI_FLASK_SECRET", "change-this-ai-flask-secret")
app.config.update(
    SESSION_COOKIE_NAME='belnipiai_session',
    SESSION_COOKIE_PATH='/',
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE='Lax'
)

MODEL_DEFAULT = os.environ.get("MODEL_DEFAULT", "gemma3:4b")
AI_SSO_SHARED_SECRET = os.environ.get("AI_SSO_SHARED_SECRET", "change-this-ai-sso-secret")
AI_SSO_MAX_AGE_SEC = int(os.environ.get("AI_SSO_MAX_AGE_SEC", "300"))

OLLAMA_HOST = os.environ.get("OLLAMA_HOST", "http://localhost:11434")
DATABASE_URL = os.environ.get("DATABASE_URL", "postgresql://postgres:sdd05072008sdd@localhost:5432/belnipiai")
KB_FULL_ANALYSIS_MAX_CHUNKS = int(os.environ.get("KB_FULL_ANALYSIS_MAX_CHUNKS", "2000"))

_client = ollama.Client(host=OLLAMA_HOST)


def _current_user() -> str:
    return _normalize_username(session.get("username") or "")


def _normalize_username(raw_username: str) -> str:
    login = (raw_username or "").strip().lower()
    if "\\" in login:
        login = login.split("\\")[-1].strip()
    if "@" in login:
        login = login.split("@")[0].strip()
    return login


def _verify_sso_payload(username: str, display_name: str, ts_raw: str, sig: str) -> bool:
    try:
        ts_val = int(ts_raw)
    except Exception:
        return False
    if abs(int(time.time()) - ts_val) > AI_SSO_MAX_AGE_SEC:
        return False
    payload = f"{username}|{display_name}|{ts_raw}"
    expected = hmac.new(
        AI_SSO_SHARED_SECRET.encode("utf-8"),
        payload.encode("utf-8"),
        hashlib.sha256
    ).hexdigest()
    return hmac.compare_digest(expected, sig or "")


# ── модели ────────────────────────────────────────────────────────────────────

# Отображаемые названия для фронтенда
MODEL_LABELS = {
    "gemma3:4b":  "Базовая",   # ноутбук
    "gemma4:e4b": "Базовая",   # сервер
    "gemma4:26b": "Pro",        # сервер
}

MODEL_PROFILES = {
    "gemma3:4b":  {"num_ctx": 16000, "chunk_tokens": 12000, "no_think": False},
    "gemma4:e4b": {"num_ctx": 32000, "chunk_tokens": 26000, "no_think": False},
    "gemma4:26b": {"num_ctx": 32000, "chunk_tokens": 26000, "no_think": False},
    "qwen3:14b":  {"num_ctx": 32000, "chunk_tokens": 26000, "no_think": False},
    "qwen3":      {"num_ctx": 32000, "chunk_tokens": 26000, "no_think": False},
}

DEFAULT_PROFILE = {"num_ctx": 32000, "chunk_tokens": 26000, "no_think": False}

MODEL_TASK_PRIORITY = {
    "image":    (["gemma4:26b", "gemma4:e4b", "gemma3:4b"], "Изображение — нужна vision-модель"),
    "document": (["gemma4:26b", "gemma4:e4b", "gemma3:4b"], "Документ — нужен большой контекст"),
    "table":    (["gemma4:26b", "gemma4:e4b", "gemma3:4b"], "Таблица — нужна аналитическая модель"),
    "complex":  (["gemma4:26b", "gemma4:e4b", "gemma3:4b"], "Сложный вопрос — нужна умная модель"),
    "simple":   (["gemma3:4b", "gemma4:e4b", "gemma4:26b"], "Простой вопрос"),
}

COMPLEX_KEYWORDS = {
    "проанализируй", "сравни", "объясни", "почему", "как работает",
    "разбери", "оцени", "составь", "напиши", "рассчитай", "найди противоречия",
}


def suggest_model(file_ext: str, question: str, available: list) -> tuple[str, str]:
    """Возвращает (model_name, reason) из доступных моделей."""

    def pick(task_key: str) -> str | None:
        preferred, _ = MODEL_TASK_PRIORITY[task_key]
        for p in preferred:
            base = p.split(":")[0]
            for m in available:
                if m == p or m.startswith(base):
                    return m
        return None

    images = {".jpg", ".jpeg", ".png", ".bmp", ".gif", ".webp", ".tiff", ".tif"}
    docs   = {".pdf", ".docx", ".doc"}
    tables = {".xlsx", ".xls", ".xlsm", ".ods", ".csv", ".tsv"}

    if file_ext in images:
        task = "image"
    elif file_ext in docs:
        task = "document"
    elif file_ext in tables:
        task = "table"
    elif len(question) > 180 or any(kw in question.lower() for kw in COMPLEX_KEYWORDS):
        task = "complex"
    else:
        task = "simple"

    model = pick(task) or (available[0] if available else MODEL_DEFAULT)
    _, reason = MODEL_TASK_PRIORITY[task]
    return model, reason


SYSTEM_PROMPT_DOC = (
    "You are a precise document analyst. "
    "Answer ONLY based on the provided document. "
    "If the answer is not found in the document, say so explicitly — never invent facts. "
    "Format your response in Markdown."
)

SYSTEM_PROMPT_CHAT = (
    "You are a helpful assistant. "
    "If you are not certain about something, say so — never make up facts, numbers or dates. "
    "Format your responses in Markdown."
)

SYSTEM_PROMPT_RAG = (
    "You are an expert assistant working with enterprise documents. "
    "Answer based on the provided context from the knowledge base. "
    "Be concise: match the answer length to the question — simple questions get short answers, complex ones get detailed answers. "
    "Use inline citations like [Источник N] only where needed — do NOT add a sources list or bibliography at the end. "
    "If no relevant information is found, say: 'В базе знаний не найдено информации по данному вопросу.' "
    "Never invent facts. Answer in the same language as the question. Format your response in Markdown."
)

SYSTEM_PROMPT_COMPARE = (
    "Ты — эксперт по проверке документов. Сравни каждую КОПИЮ с ОРИГИНАЛОМ.\n\n"
    "Данные представлены абзацами: сначала полный текст абзаца, затем его параметры форматирования.\n\n"
    "НАЙДИ ВСЕ РАСХОЖДЕНИЯ двух видов:\n\n"
    "### 1. ТЕКСТ (самое важное)\n"
    "- Удалённый текст: абзацы или фрагменты, которые есть в оригинале, но отсутствуют в копии\n"
    "- Добавленный текст: есть в копии, нет в оригинале\n"
    "- Изменённый текст: слова заменены, переставлены, опечатки, другая формулировка\n"
    "- Структурные изменения: абзацы объединены или разбиты\n\n"
    "### 2. ФОРМАТИРОВАНИЕ\n"
    "- Поля страницы (верх/низ/лево/право в см)\n"
    "- Шрифт: название, размер (pt), начертание (жирный/курсив/подчёркнутый)\n"
    "- Отступы: красная строка, левый отступ (в см)\n"
    "- Выравнивание: слева / по центру / справа / по ширине\n"
    "- Межстрочный интервал и интервалы до/после абзаца\n\n"
    "ФОРМАТ ОТВЕТА для каждой копии:\n"
    "## Копия N — «имя файла»\n"
    "### Текстовые расхождения\n"
    "- **[Тип]** Абзац X: оригинал → «...» / копия → «...»\n"
    "### Расхождения форматирования\n"
    "- **[Параметр]** Абзац X: оригинал → копия\n\n"
    "Если расхождений нет — пиши: '✅ Документ полностью соответствует оригиналу'.\n"
    "Будь точным: цитируй реальный текст из данных. Не придумывай."
)


def get_profile(model: str) -> dict:
    if model in MODEL_PROFILES:
        return MODEL_PROFILES[model]
    for key, profile in MODEL_PROFILES.items():
        if model.startswith(key.split(":")[0]):
            return profile
    return DEFAULT_PROFILE


def apply_no_think(text: str, model: str) -> str:
    if get_profile(model)["no_think"]:
        return "/no_think " + text
    return text


def estimate_tokens(text: str) -> int:
    """Быстрая локальная оценка без сетевого вызова (для решений о чанкинге)."""
    cyrillic = sum(1 for c in text if 'Ѐ' <= c <= 'ӿ')
    return max(1, cyrillic // 2 + (len(text) - cyrillic) // 3)


def count_tokens(text: str, model: str = MODEL_DEFAULT) -> int:
    """Точный подсчёт через Ollama /api/tokenize, fallback на estimate_tokens."""
    try:
        body = json.dumps({"model": model, "prompt": text}).encode()
        req = urllib.request.Request(
            f"{OLLAMA_HOST}/api/tokenize",
            data=body,
            headers={"Content-Type": "application/json"},
            method="POST"
        )
        with urllib.request.urlopen(req, timeout=3) as resp:
            return max(1, len(json.loads(resp.read()).get("tokens", [])))
    except Exception:
        return estimate_tokens(text)


# ── база данных ───────────────────────────────────────────────────────────────

def init_db():
    schema = os.path.join(ROOT_DIR, "schema.sql")
    if not os.path.exists(schema):
        return
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(open(schema).read())
        conn.commit()


# ── управление контекстом ─────────────────────────────────────────────────────

def load_history(conv_id: str) -> list:
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT role, content FROM messages "
                "WHERE conversation_id = %s AND role != 'summary' "
                "ORDER BY created_at",
                (conv_id,)
            )
            rows = cur.fetchall()
    return [{"role": r["role"], "content": r["content"]} for r in rows]


def load_history_with_summary(conv_id: str, num_ctx: int) -> list:
    """Загружает историю. Для сообщений с файлами восстанавливает документ из БД."""
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT content, created_at FROM messages "
                "WHERE conversation_id = %s AND role = 'summary' "
                "ORDER BY created_at DESC LIMIT 1",
                (conv_id,)
            )
            summary_row = cur.fetchone()

            since = summary_row["created_at"] if summary_row else None
            base_q = """
                SELECT m.id, m.role, m.content, m.token_count,
                       a.extracted_text, a.original_name
                FROM messages m
                LEFT JOIN attachments a ON a.message_id = m.id
                WHERE m.conversation_id = %s AND m.role != 'summary'
            """
            if since:
                cur.execute(base_q + " AND m.created_at > %s ORDER BY m.created_at", (conv_id, since))
            else:
                cur.execute(base_q + " ORDER BY m.created_at", (conv_id,))
            rows = cur.fetchall()

    rows = list(rows)

    # Только последний документ получает полный текст в контекст,
    # старые документы — лишь ссылка на имя файла
    max_doc_chars = int(num_ctx * 0.55) * 4
    last_doc_idx = max(
        (i for i, r in enumerate(rows) if r["role"] == "user" and r["extracted_text"]),
        default=None
    )

    def build_content(row, idx):
        if row["role"] == "user" and row["extracted_text"]:
            if idx == last_doc_idx:
                doc = row["extracted_text"]
                if len(doc) > max_doc_chars:
                    doc = doc[:max_doc_chars] + "\n\n[... документ обрезан из-за лимита контекста ...]"
                return f"<document name=\"{row['original_name']}\">\n{doc}\n</document>\n\nВопрос: {row['content']}"
            return f"[Ранее загружен документ: {row['original_name']}]\n\nВопрос: {row['content']}"
        return row["content"]

    messages = []
    if summary_row:
        messages.append({"role": "user", "content": f"[Резюме предыдущего разговора]\n{summary_row['content']}"})
        messages.append({"role": "assistant", "content": "Понял, продолжаю разговор с учётом предыдущего контекста."})

    # Считаем токены по фактическому контенту (с документом), а не по r["content"].
    # Это правильно при смене модели: сохранённый token_count от старой модели не используется.
    total_tokens = sum(estimate_tokens(build_content(r, i)) for i, r in enumerate(rows))
    threshold = int(num_ctx * 0.85)

    if total_tokens <= threshold or len(rows) <= 4:
        messages += [{"role": r["role"], "content": build_content(r, i)} for i, r in enumerate(rows)]
        return messages

    # Нужна обрезка. Бюджет = threshold минус уже добавленные summary-сообщения.
    budget = threshold - sum(estimate_tokens(m["content"]) for m in messages)

    # Якоря: последний документ + ответ на него — всегда включаем, независимо от бюджета.
    anchors: set[int] = set()
    if last_doc_idx is not None:
        anchors.add(last_doc_idx)
        nxt = last_doc_idx + 1
        if nxt < len(rows) and rows[nxt]["role"] == "assistant":
            anchors.add(nxt)

    for idx in sorted(anchors):
        budget -= estimate_tokens(build_content(rows[idx], idx))

    # Добираем свежие сообщения (от конца) в рамках оставшегося бюджета.
    tail: set[int] = set()
    for i in range(len(rows) - 1, -1, -1):
        if i in anchors:
            continue
        t = estimate_tokens(rows[i]["content"])
        if budget - t < 100:
            break
        tail.add(i)
        budget -= t

    keep = sorted(anchors | tail)
    messages += [{"role": rows[i]["role"], "content": build_content(rows[i], i)} for i in keep]

    return messages


def compact_history_for_rag(history: list, max_messages: int = 8,
                            max_chars: int = 1600) -> list:
    compact = []
    for msg in history[-max_messages:]:
        content = msg.get("content", "")
        if "<document name=" in content:
            content = re.sub(
                r"<document name=\"([^\"]+)\">.*?</document>",
                r"[Документ в истории: \1]",
                content,
                flags=re.S,
            )
        if len(content) > max_chars:
            content = content[:max_chars] + "\n\n[... история обрезана ...]"
        compact.append({"role": msg.get("role", "user"), "content": content})
    return compact


def is_full_kb_analysis(question: str) -> bool:
    q = question.lower()
    broad_markers = (
        "все файлы", "все документы", "всю базу", "базу знаний",
        "полностью", "целиком", "по всем файлам", "по всем документам",
        "проанализируй все", "анализ всех", "сравни все",
    )
    analysis_markers = (
        "проанализ", "анализ", "сравни", "сводк", "вывод",
        "найди", "проверь", "обобщ", "резюм",
    )
    return any(m in q for m in broad_markers) and any(m in q for m in analysis_markers)


def maybe_summarize(conv_id: str, model: str, num_ctx: int):
    """Если история занимает > 65% контекста — делаем резюме (вызывается из фонового треда)."""
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT SUM(token_count) FROM messages "
                "WHERE conversation_id = %s AND role != 'summary'",
                (conv_id,)
            )
            total = cur.fetchone()["sum"] or 0

    if total < int(num_ctx * 0.65):
        return

    history = load_history(conv_id)
    if len(history) < 6:
        return

    old_messages = history[:-4]
    text_to_summarize = "\n".join(f"{m['role'].upper()}: {m['content']}" for m in old_messages)

    resp = _client.chat(
        model=model,
        messages=[{
            "role": "user",
            "content": f"Сделай краткое резюме этого диалога (важные факты, решения, контекст):\n\n{text_to_summarize}"
        }],
        options={"num_ctx": num_ctx, "temperature": 0.1},
        stream=False
    )
    summary_text = resp["message"]["content"]

    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO messages (conversation_id, role, content, token_count) VALUES (%s, 'summary', %s, %s)",
                (conv_id, summary_text, estimate_tokens(summary_text))
            )
        conn.commit()


def save_message(conv_id: str, role: str, content: str, model: str = MODEL_DEFAULT,
                 attachment_name: str = None, attachment_text: str = None):
    tokens = estimate_tokens(content)
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO messages (id, conversation_id, role, content, token_count) VALUES (%s, %s, %s, %s, %s) RETURNING id",
                (str(uuid.uuid4()), conv_id, role, content, tokens)
            )
            msg_id = cur.fetchone()["id"]
            if attachment_name and attachment_text:
                cur.execute(
                    "INSERT INTO attachments (message_id, original_name, extracted_text) VALUES (%s, %s, %s)",
                    (msg_id, attachment_name, attachment_text)
                )
            cur.execute(
                "UPDATE conversations SET updated_at = NOW() WHERE id = %s",
                (conv_id,)
            )
        conn.commit()


def generate_title(question: str, model: str) -> str:
    try:
        resp = _client.chat(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Ты генерируешь названия чатов. "
                        "Отвечай ТОЛЬКО самим названием — одной строкой, 2–4 слова. "
                        "Никаких объяснений, никаких списков, никаких кавычек."
                    )
                },
                {
                    "role": "user",
                    "content": f"Название чата по запросу:\n{question[:400]}"
                }
            ],
            options={"num_ctx": 2048, "temperature": 0.1},
            stream=False
        )
        raw = resp["message"]["content"]
        # Убираем теги <think>...</think>
        raw = re.sub(r"<think>.*?</think>", "", raw, flags=re.DOTALL).strip()
        # Если модель вернула список — берём первый пункт
        m = re.search(r"^[\d\-\*•]\s*\.?\s*(.+)$", raw, re.MULTILINE)
        title = m.group(1).strip() if m else next(
            (ln.strip() for ln in raw.splitlines() if ln.strip()), ""
        )
        # Убираем типичные «мусорные» префиксы
        title = re.sub(
            r"(?i)^(название|title|тема|чат|вот|итого|ответ)\s*[:\-–—]?\s*", "", title
        ).strip()
        title = title.strip('"\'«»').rstrip(".,;:!?").strip()
        return title[:60] if len(title) > 2 else question[:50]
    except Exception:
        return question[:50]


# ── утилиты файлов ────────────────────────────────────────────────────────────

def to_base64(pil_image):
    buf = io.BytesIO()
    pil_image.save(buf, format="PNG", optimize=False, compress_level=0)
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def get_available_models():
    try:
        result = _client.list()
        return [m.model for m in result.models]
    except Exception:
        return [MODEL_DEFAULT]


def read_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def detect_pdf_type(path):
    import pymupdf
    doc = pymupdf.open(path)
    total_chars = sum(len(p.get_text().strip()) for p in doc)
    avg = total_chars / max(len(doc), 1)
    return "scanned" if avg < 50 else "text"


def read_pdf_text(path):
    import pymupdf
    doc = pymupdf.open(path)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            pages.append(f"[Страница {i+1}]\n{text}")
    return pages


def read_scanned_pdf_as_text(path: str, model: str, profile: dict) -> str:
    from pdf2image import convert_from_path
    pil_pages = convert_from_path(path, dpi=200)
    page_texts = []
    for i, pil_page in enumerate(pil_pages):
        b64 = to_base64(pil_page)
        r = _client.chat(
            model=model,
            messages=[{"role": "user",
                       "content": f"Извлеки весь текст со страницы {i+1}.",
                       "images": [b64]}],
            options={"num_ctx": profile["num_ctx"], "temperature": 0.0},
            stream=False,
        )
        page_texts.append(f"[Страница {i+1}]\n{r['message']['content']}")
    return "\n\n".join(page_texts)


def read_docx(path):
    import docx
    doc = docx.Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def read_excel(path):
    import pandas as pd
    ext = os.path.splitext(path)[1].lower()
    parts = []
    with pd.ExcelFile(path, engine="xlrd" if ext == ".xls" else "openpyxl") as xf:
        for sheet in xf.sheet_names:
            df = xf.parse(sheet)
            parts.append(f"[Лист: {sheet}]\n{df.to_string()}")
    return "\n\n".join(parts)


def read_excel_sheets(path: str) -> list[dict]:
    """Return structured sheet data for RAG chunking (not flattened to string)."""
    import pandas as pd

    def is_empty(v) -> bool:
        try:
            if pd.isna(v):
                return True
        except Exception:
            pass
        return str(v).strip() == ""

    def clean_cell(v) -> str:
        return "" if is_empty(v) else str(v).strip()

    def dedupe_headers(headers: list[str]) -> list[str]:
        seen = {}
        out = []
        for i, h in enumerate(headers, 1):
            base = h.strip() or f"Колонка {i}"
            seen[base] = seen.get(base, 0) + 1
            out.append(base if seen[base] == 1 else f"{base} {seen[base]}")
        return out

    def infer_header_row(df) -> int | None:
        candidates = list(df.index[: min(len(df.index), 25)])
        best_idx = None
        best_score = -1.0
        for idx in candidates:
            vals = [clean_cell(v) for v in df.loc[idx].tolist()]
            non_empty = [v for v in vals if v]
            if not non_empty:
                continue
            text_like = sum(1 for v in non_empty if any(ch.isalpha() for ch in v))
            unique = len(set(non_empty))
            numeric_like = sum(1 for v in non_empty if v.replace(".", "", 1).isdigit())
            next_rows = df.loc[df.index[df.index.get_loc(idx) + 1: df.index.get_loc(idx) + 4]]
            has_data_after = not next_rows.dropna(how="all").empty
            score = len(non_empty) + text_like * 0.7 + unique * 0.2 - numeric_like * 0.35
            if has_data_after:
                score += 1.5
            if score > best_score:
                best_score = score
                best_idx = idx
        return best_idx

    ext = os.path.splitext(path)[1].lower()
    if ext == ".xls":
        engine = "xlrd"
    elif ext == ".ods":
        engine = "odf"
    else:
        engine = "openpyxl"
    sheets = []
    with pd.ExcelFile(path, engine=engine) as xf:
        for sheet_name in xf.sheet_names:
            raw = xf.parse(sheet_name, header=None, dtype=object)
            raw = raw.dropna(how="all").dropna(axis=1, how="all")
            if raw.empty:
                continue
            header_idx = infer_header_row(raw)
            if header_idx is None:
                headers = [f"Колонка {i+1}" for i in range(raw.shape[1])]
                data = raw
            else:
                headers = dedupe_headers([clean_cell(v) for v in raw.loc[header_idx].tolist()])
                data = raw.loc[raw.index[raw.index.get_loc(header_idx) + 1:]]
            data = data.dropna(how="all")
            rows = []
            for idx, row in data.iterrows():
                values = row.tolist()
                if all(is_empty(v) for v in values):
                    continue
                rows.append({
                    "row_number": int(idx) + 1,
                    "values": values,
                })
            if not rows:
                continue
            sheets.append({
                "name":    sheet_name,
                "headers": headers,
                "rows":    rows,
            })
    return sheets


def read_csv(path, ext=".csv"):
    import pandas as pd
    sep = "\t" if ext == ".tsv" else ","
    try:
        df = pd.read_csv(path, sep=sep, encoding="utf-8", errors="ignore")
    except Exception:
        df = pd.read_csv(path, sep=sep, encoding="cp1251", errors="ignore")
    return df.to_string()


# ── стриминг ──────────────────────────────────────────────────────────────────

def stream_text_response(content, question, model, history=None):
    profile = get_profile(model)
    num_ctx = profile["num_ctx"]

    max_chars = (num_ctx - 1000) * 4
    if len(content) > max_chars:
        content = content[:max_chars]

    messages = [{"role": "system", "content": SYSTEM_PROMPT_DOC}]
    if history:
        messages += history
    messages.append({
        "role": "user",
        "content": apply_no_think(f"<document>\n{content}\n</document>\n\nВопрос: {question}", model)
    })

    response = _client.chat(model=model, messages=messages,
                            options={"num_ctx": num_ctx, "temperature": 0.1}, stream=True)
    for chunk in response:
        yield f"data: {json.dumps({'text': chunk['message']['content']})}\n\n"
    yield "data: [DONE]\n\n"


def stream_image_response(images_b64, question, model, history=None):
    profile = get_profile(model)
    messages = []
    if history:
        messages += history
    messages.append({
        "role": "user",
        "content": apply_no_think(question, model),
        "images": images_b64
    })
    response = _client.chat(model=model, messages=messages,
                            options={"num_ctx": profile["num_ctx"], "temperature": 0.1}, stream=True)
    for chunk in response:
        yield f"data: {json.dumps({'text': chunk['message']['content']})}\n\n"
    yield "data: [DONE]\n\n"


def stream_map_reduce(pages_text, question, model, history=None):
    profile = get_profile(model)
    num_ctx = profile["num_ctx"]
    chunk_chars = profile["chunk_tokens"] * 4

    chunks, current, cur_len = [], [], 0
    for page in pages_text:
        if cur_len + len(page) > chunk_chars and current:
            chunks.append("\n\n".join(current))
            current, cur_len = [page], len(page)
        else:
            current.append(page)
            cur_len += len(page)
    if current:
        chunks.append("\n\n".join(current))

    total = len(chunks)

    if total == 1:
        yield from stream_text_response(chunks[0], question, model, history)
        return

    yield f"data: {json.dumps({'text': f'📦 Документ большой — обрабатываю {total} частей параллельно...\\n\\n'})}\n\n"

    def process_chunk(i, chunk):
        prompt = apply_no_think(
            f"Это часть {i+1} из {total} документа.\nВопрос: {question}\n"
            f"Извлеки всё относящееся к вопросу, сохрани детали, цифры, имена:\n\n{chunk}", model
        )
        s = ""
        for r in _client.chat(model=model,
                               messages=[{"role": "user", "content": prompt}],
                               options={"num_ctx": num_ctx, "temperature": 0.1}, stream=True):
            s += r["message"]["content"]
        return i, f"[Часть {i+1}/{total}]\n{s}"

    summaries = [None] * total
    with ThreadPoolExecutor(max_workers=min(total, 3)) as pool:
        futures = {pool.submit(process_chunk, i, chunk): i for i, chunk in enumerate(chunks)}
        for future in as_completed(futures):
            i, result = future.result()
            summaries[i] = result
            done = sum(1 for s in summaries if s is not None)
            yield f"data: {json.dumps({'text': f'✅ Часть {i+1} готова ({done}/{total})\\n'})}\n\n"

    yield f"data: {json.dumps({'text': '\\n---\\n📝 **Финальный ответ**\\n\\n'})}\n\n"
    yield from stream_text_response("\n\n".join(summaries), question, model, history)


# ── роуты: конверсации ────────────────────────────────────────────────────────

@app.route("/api/conversations", methods=["GET"])
def api_list_conversations():
    username = _current_user()
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id, title, model, updated_at FROM conversations "
                "WHERE username = %s ORDER BY updated_at DESC",
                (username,)
            )
            rows = cur.fetchall()
    return {"conversations": [dict(r) for r in rows]}


@app.route("/api/conversations", methods=["POST"])
def api_create_conversation():
    data = request.get_json(force=True)
    model = data.get("model", MODEL_DEFAULT)
    cid = str(uuid.uuid4())
    username = _current_user()
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO conversations (id, username, title, model) VALUES (%s, %s, %s, %s)",
                (cid, username, "Новый чат", model)
            )
        conn.commit()
    return {"id": cid, "title": "Новый чат", "model": model}


@app.route("/api/conversations/<conv_id>", methods=["DELETE"])
def api_delete_conversation(conv_id):
    username = _current_user()
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "DELETE FROM conversations WHERE id = %s AND username = %s",
                (conv_id, username)
            )
        conn.commit()
    return {"ok": True}


@app.route("/api/conversations/<conv_id>/messages", methods=["GET"])
def api_get_messages(conv_id):
    username = _current_user()
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT c.id FROM conversations c WHERE c.id = %s AND c.username = %s",
                (conv_id, username)
            )
            if not cur.fetchone():
                return {"messages": []}
            cur.execute(
                "SELECT m.id, m.role, m.content, m.created_at, "
                "a.original_name as file_name "
                "FROM messages m "
                "LEFT JOIN attachments a ON a.message_id = m.id "
                "WHERE m.conversation_id = %s AND m.role != 'summary' "
                "ORDER BY m.created_at",
                (conv_id,)
            )
            rows = cur.fetchall()
    return {"messages": [dict(r) for r in rows]}


@app.route("/api/conversations/<conv_id>/title", methods=["PATCH"])
def api_update_title(conv_id):
    data = request.get_json(force=True)
    title = data.get("title", "")[:80]
    username = _current_user()
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE conversations SET title = %s WHERE id = %s AND username = %s",
                (title, conv_id, username)
            )
        conn.commit()
    return {"ok": True}


# ── роут: чат ─────────────────────────────────────────────────────────────────

@app.route("/api/conversations/<conv_id>/chat", methods=["POST"])
def api_chat(conv_id):
    question  = request.form.get("question", "").strip()
    model     = request.form.get("model", MODEL_DEFAULT)
    file      = request.files.get("file")
    use_rag   = request.form.get("use_rag", "false").lower() == "true"
    folder_id = request.form.get("folder_id") or None

    if not question:
        return {"error": "Вопрос не может быть пустым"}, 400

    username = _current_user()
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id FROM conversations WHERE id = %s AND username = %s",
                (conv_id, username)
            )
            if not cur.fetchone():
                return {"error": "Чат не найден"}, 404

    profile = get_profile(model)

    # Сохраняем файл до генератора
    tmp_path = file_name = extracted_text = None
    ext = ""
    if file and file.filename:
        ext = os.path.splitext(file.filename)[1].lower()
        file_name = file.filename
        tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
        tmp_path = tmp.name
        tmp.close()
        file.save(tmp_path)

    # Загружаем историю для контекста
    history = load_history_with_summary(conv_id, profile["num_ctx"])

    # Сохраняем вопрос пользователя
    save_message(conv_id, "user", question, model)

    doc_request = is_document_request(question)
    doc_fmt     = detect_format(question) if doc_request else "docx"

    def generate():
        nonlocal extracted_text
        full_response = []

        try:
            if not tmp_path:
                if use_rag:
                    rag_history = compact_history_for_rag(history)
                    if is_full_kb_analysis(question):
                        pages, total_chunks = fetch_all_context_pages(
                            folder_id=folder_id,
                            max_chunks=KB_FULL_ANALYSIS_MAX_CHUNKS,
                        )
                        if pages:
                            if total_chunks > len(pages):
                                yield f"data: {json.dumps({'text': f'База знаний очень большая: обрабатываю первые {len(pages)} из {total_chunks} чанков. Увеличить лимит можно через KB_FULL_ANALYSIS_MAX_CHUNKS.\\n\\n'})}\n\n"
                            gen = stream_map_reduce(pages, question, model, rag_history)
                            for chunk in gen:
                                if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                                    try:
                                        full_response.append(json.loads(chunk[6:]).get("text", ""))
                                    except Exception:
                                        pass
                                yield chunk
                            return

                    reserved_tokens = (
                        estimate_tokens(question)
                        + sum(estimate_tokens(m["content"]) for m in rag_history)
                        + 4500
                    )
                    rag_budget = max(6000, int(profile["num_ctx"] * 0.90) - reserved_tokens)
                    rag_chunks = build_context_chunks(
                        question,
                        folder_id=folder_id,
                        max_tokens=rag_budget,
                        include_coverage=is_full_kb_analysis(question),
                    )
                    if rag_chunks:
                        rag_context = build_rag_context(rag_chunks, max_tokens=rag_budget)
                        sources = [
                            {
                                "name":  c["original_name"],
                                "page":  c.get("page_num"),
                                "score": round(float(c["score"]), 3),
                            }
                            for c in rag_chunks
                        ]
                        yield f"data: {json.dumps({'sources': sources})}\n\n"
                        messages = [{"role": "system", "content": SYSTEM_PROMPT_RAG}]
                        messages += rag_history
                        messages.append({
                            "role": "user",
                            "content": (
                                f"<context>\n{rag_context}\n</context>\n\n"
                                f"Вопрос: {question}"
                            ),
                        })
                        temp = 0.1
                    else:
                        yield f"data: {json.dumps({'text': '_В базе знаний не найдено релевантных документов. Отвечаю из общих знаний._\\n\\n'})}\n\n"
                        messages = [{"role": "system", "content": SYSTEM_PROMPT_CHAT}]
                        messages += history
                        messages.append({"role": "user", "content": apply_no_think(question, model)})
                        temp = 0.7
                else:
                    if doc_request:
                        messages = [{"role": "system", "content": SYSTEM_PROMPT_DOCGEN}]
                        messages += history
                        messages.append({"role": "user", "content": apply_no_think(question, model)})
                        temp = 0.3
                    else:
                        messages = [{"role": "system", "content": SYSTEM_PROMPT_CHAT}]
                        messages += history
                        messages.append({"role": "user", "content": apply_no_think(question, model)})
                        temp = 0.7

                for chunk in _client.chat(model=model, messages=messages,
                                          options={"num_ctx": profile["num_ctx"], "temperature": temp},
                                          stream=True):
                    delta = chunk["message"]["content"]
                    full_response.append(delta)
                    yield f"data: {json.dumps({'text': delta})}\n\n"
                yield "data: [DONE]\n\n"

            else:
                if ext in (".txt", ".md", ".log"):
                    text = read_txt(tmp_path)
                    extracted_text = text
                    gen = (stream_text_response(text, question, model, history)
                           if estimate_tokens(text) <= profile["chunk_tokens"]
                           else stream_map_reduce(
                               [text[i:i+profile["chunk_tokens"]*4]
                                for i in range(0, len(text), profile["chunk_tokens"]*4)],
                               question, model, history))
                    for chunk in gen:
                        if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                            try:
                                full_response.append(json.loads(chunk[6:]).get("text", ""))
                            except Exception:
                                pass
                        yield chunk

                elif ext == ".pdf":
                    pdf_type = detect_pdf_type(tmp_path)
                    if pdf_type == "text":
                        pages = read_pdf_text(tmp_path)
                        extracted_text = "\n\n".join(pages)
                        gen = (stream_text_response(extracted_text, question, model, history)
                               if estimate_tokens(extracted_text) <= profile["chunk_tokens"]
                               else stream_map_reduce(pages, question, model, history))
                        for chunk in gen:
                            if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                                try:
                                    full_response.append(json.loads(chunk[6:]).get("text", ""))
                                except Exception:
                                    pass
                            yield chunk
                    else:
                        yield f"data: {json.dumps({'text': '🖼️ Сканированный PDF — конвертирую...\\n\\n'})}\n\n"
                        from pdf2image import convert_from_path
                        pil_pages = convert_from_path(tmp_path, dpi=200)
                        if len(pil_pages) <= 5:
                            images_b64 = [to_base64(p) for p in pil_pages]
                            for chunk in stream_image_response(images_b64, question, model, history):
                                if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                                    try:
                                        full_response.append(json.loads(chunk[6:]).get("text", ""))
                                    except Exception:
                                        pass
                                yield chunk
                        else:
                            n_pages = len(pil_pages)
                            yield f"data: {json.dumps({'text': f'🔍 Извлекаю текст из {n_pages} страниц параллельно...\\n\\n'})}\n\n"

                            def extract_page(args):
                                idx, pil_page = args
                                b64 = to_base64(pil_page)
                                r = _client.chat(
                                    model=model,
                                    messages=[{"role": "user",
                                               "content": apply_no_think(f"Страница {idx+1} из {n_pages}. Извлеки весь текст.", model),
                                               "images": [b64]}],
                                    options={"num_ctx": profile["num_ctx"], "temperature": 0.1},
                                    stream=False
                                )
                                return idx, f"[Страница {idx+1}]\n{r['message']['content']}"

                            page_dict = {}
                            with ThreadPoolExecutor(max_workers=min(n_pages, 3)) as pool:
                                futs = {pool.submit(extract_page, (i, p)): i for i, p in enumerate(pil_pages)}
                                for fut in as_completed(futs):
                                    idx, text = fut.result()
                                    page_dict[idx] = text
                                    yield f"data: {json.dumps({'text': f'✅ Страница {idx+1}/{n_pages} готова\\n'})}\n\n"
                            page_texts = [page_dict[i] for i in range(n_pages)]
                            extracted_text = "\n\n".join(page_texts)
                            for chunk in stream_map_reduce(page_texts, question, model, history):
                                if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                                    try:
                                        full_response.append(json.loads(chunk[6:]).get("text", ""))
                                    except Exception:
                                        pass
                                yield chunk

                elif ext in (".docx", ".doc"):
                    text = read_docx(tmp_path)
                    extracted_text = text
                    gen = (stream_text_response(text, question, model, history)
                           if estimate_tokens(text) <= profile["chunk_tokens"]
                           else stream_map_reduce(
                               [text[i:i+profile["chunk_tokens"]*4]
                                for i in range(0, len(text), profile["chunk_tokens"]*4)],
                               question, model, history))
                    for chunk in gen:
                        if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                            try:
                                full_response.append(json.loads(chunk[6:]).get("text", ""))
                            except Exception:
                                pass
                        yield chunk

                elif ext in (".xlsx", ".xls", ".xlsm", ".ods"):
                    text = read_excel(tmp_path)
                    extracted_text = text
                    for chunk in stream_text_response(text, question, model, history):
                        if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                            try:
                                full_response.append(json.loads(chunk[6:]).get("text", ""))
                            except Exception:
                                pass
                        yield chunk

                elif ext in (".csv", ".tsv"):
                    text = read_csv(tmp_path, ext)
                    extracted_text = text
                    for chunk in stream_text_response(text, question, model, history):
                        if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                            try:
                                full_response.append(json.loads(chunk[6:]).get("text", ""))
                            except Exception:
                                pass
                        yield chunk

                elif ext in (".jpg", ".jpeg", ".png", ".bmp", ".gif", ".webp", ".tiff", ".tif"):
                    with open(tmp_path, "rb") as f:
                        b64 = base64.b64encode(f.read()).decode("utf-8")
                    for chunk in stream_image_response([b64], question, model, history):
                        if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                            try:
                                full_response.append(json.loads(chunk[6:]).get("text", ""))
                            except Exception:
                                pass
                        yield chunk
                else:
                    yield f"data: {json.dumps({'error': f'Формат {ext} не поддерживается'})}\n\n"

        finally:
            if tmp_path:
                os.unlink(tmp_path)

            # Сохраняем ответ ассистента
            assistant_text = "".join(full_response)
            if assistant_text:
                save_message(conv_id, "assistant", assistant_text, model,
                             attachment_name=file_name,
                             attachment_text=extracted_text)

                # Генерация документа
                if doc_request:
                    code = extract_code(assistant_text)

                    # Если AI не сгенерировал код — принудительный повторный запрос
                    if not code:
                        retry_messages = [
                            {"role": "system", "content": SYSTEM_PROMPT_DOCGEN},
                            {"role": "user",   "content": question},
                            {"role": "assistant", "content": assistant_text},
                            {"role": "user",   "content": (
                                f"Ты не создал файл. Напиши ТОЛЬКО Python-код "
                                f"в блоке ```python:document для создания {doc_fmt}-файла "
                                f"с данными из предыдущего ответа. Никакого текста — только код."
                            )},
                        ]
                        retry_resp = _client.chat(
                            model=model,
                            messages=retry_messages,
                            options={"num_ctx": profile["num_ctx"], "temperature": 0.2},
                            stream=False,
                        )
                        code = extract_code(retry_resp["message"]["content"])

                    if code:
                        out_path, out_name = make_output_path(doc_fmt)
                        ok, err = run_code(code, out_path)
                        if ok:
                            yield f"data: {json.dumps({'document': {'url': f'/api/download/{out_name}', 'name': out_name, 'fmt': doc_fmt}})}\n\n"
                        else:
                            yield f"data: {json.dumps({'text': f'\\n\\n❌ Не удалось создать файл: {err}'})}\n\n"
                    else:
                        yield f"data: {json.dumps({'text': '\\n\\n❌ Не удалось сгенерировать код файла. Попробуйте переформулировать запрос.'})}\n\n"

                # Генерируем название чата после первого ответа
                with get_db() as conn:
                    with conn.cursor() as cur:
                        cur.execute(
                            "SELECT COUNT(*) as cnt FROM messages WHERE conversation_id = %s AND role = 'user'",
                            (conv_id,)
                        )
                        cnt = cur.fetchone()["cnt"]

                if cnt == 1:
                    title = generate_title(question, model)
                    with get_db() as conn:
                        with conn.cursor() as cur:
                            cur.execute("UPDATE conversations SET title = %s WHERE id = %s", (title, conv_id))
                        conn.commit()
                    yield f"data: {json.dumps({'title': title, 'conv_id': conv_id})}\n\n"

                threading.Thread(
                    target=maybe_summarize,
                    args=(conv_id, model, profile["num_ctx"]),
                    daemon=True
                ).start()

    return Response(stream_with_context(generate()), mimetype="text/event-stream")


# ── Скачивание сгенерированных документов ─────────────────────────────────────

@app.route("/api/download/<filename>")
def api_download(filename):
    if not re.match(r'^[0-9a-f\-]{36}\.(docx|xlsx|pdf)$', filename):
        return {"error": "Недопустимое имя файла"}, 400
    return send_from_directory(DOC_TEMP_DIR, filename, as_attachment=True)


# ── База знаний: папки ────────────────────────────────────────────────────────

@app.route("/api/kb/folders", methods=["GET"])
def api_kb_folders():
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT id, name, description FROM doc_folders ORDER BY name")
            return {"folders": [dict(r) for r in cur.fetchall()]}


@app.route("/api/kb/folders", methods=["POST"])
def api_kb_create_folder():
    data = request.get_json(force=True)
    name = (data.get("name") or "").strip()
    if not name:
        return {"error": "Имя папки не может быть пустым"}, 400
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO doc_folders (name, description) VALUES (%s, %s) RETURNING id, name",
                (name, data.get("description", ""))
            )
            row = cur.fetchone()
        conn.commit()
    return dict(row), 201


@app.route("/api/kb/folders/<folder_id>", methods=["DELETE"])
def api_kb_delete_folder(folder_id):
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM doc_folders WHERE id=%s", (folder_id,))
        conn.commit()
    return {"ok": True}


# ── База знаний: документы ────────────────────────────────────────────────────

@app.route("/api/kb/documents", methods=["GET"])
def api_kb_documents():
    folder_id = request.args.get("folder_id")
    with get_db() as conn:
        with conn.cursor() as cur:
            if folder_id:
                cur.execute(
                    "SELECT id, original_name, status, chunk_count, created_at, meta "
                    "FROM kb_documents WHERE folder_id=%s ORDER BY created_at DESC",
                    (folder_id,)
                )
            else:
                cur.execute(
                    "SELECT id, original_name, status, chunk_count, created_at, meta "
                    "FROM kb_documents ORDER BY created_at DESC"
                )
            return {"documents": [dict(r) for r in cur.fetchall()]}


@app.route("/api/kb/documents/<doc_id>", methods=["GET"])
def api_kb_document_status(doc_id):
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id, original_name, status, chunk_count, meta FROM kb_documents WHERE id=%s",
                (doc_id,)
            )
            row = cur.fetchone()
    if not row:
        return {"error": "Документ не найден"}, 404
    return dict(row)


@app.route("/api/kb/upload", methods=["POST"])
def api_kb_upload():
    file      = request.files.get("file")
    folder_id = request.form.get("folder_id") or None

    if not file or not file.filename:
        return {"error": "Файл не передан"}, 400

    ext      = os.path.splitext(file.filename)[1].lower()
    tmp      = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
    tmp_path = tmp.name
    tmp.close()
    file.save(tmp_path)

    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO kb_documents (original_name, folder_id, doc_type, status) "
                "VALUES (%s, %s, %s, 'pending') RETURNING id",
                (file.filename, folder_id, ext.lstrip("."))
            )
            doc_id = cur.fetchone()["id"]
        conn.commit()

    threading.Thread(
        target=_index_document_async,
        args=(str(doc_id), tmp_path, ext, file.filename),
        daemon=True,
    ).start()

    return {"id": str(doc_id), "status": "pending"}, 202


@app.route("/api/kb/documents/<doc_id>", methods=["DELETE"])
def api_kb_delete_document(doc_id):
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM kb_documents WHERE id=%s", (doc_id,))
        conn.commit()
    invalidate_kb_caches()
    return {"ok": True}


@app.route("/api/kb/search", methods=["POST"])
def api_kb_search():
    data      = request.get_json(force=True)
    query     = (data.get("query") or "").strip()
    folder_id = data.get("folder_id")
    top_k     = int(data.get("top_k", 30))
    if not query:
        return {"results": []}
    results = hybrid_search(query, top_k=top_k, folder_id=folder_id)
    return {"results": expand_with_neighbors(results)}


# ── Фоновая индексация ────────────────────────────────────────────────────────

def _index_document_async(doc_id: str, tmp_path: str, ext: str, original_name: str = ""):
    try:
        if ext in (".txt", ".md", ".log"):
            text   = read_txt(tmp_path)
            chunks = list(rag_chunk_text(text))

        elif ext == ".pdf":
            if detect_pdf_type(tmp_path) == "text":
                pages  = read_pdf_text(tmp_path)
                chunks = list(rag_chunk_pages(pages))
            else:
                from pdf2image import convert_from_path
                pil_pages  = convert_from_path(tmp_path, dpi=200)
                page_texts = []
                for i, pil_page in enumerate(pil_pages):
                    b64 = to_base64(pil_page)
                    r   = _client.chat(
                        model=MODEL_DEFAULT,
                        messages=[{"role": "user",
                                   "content": f"Извлеки весь текст со страницы {i+1}.",
                                   "images": [b64]}],
                        options={"num_ctx": 4096, "temperature": 0.0},
                        stream=False,
                    )
                    page_texts.append(r["message"]["content"])
                chunks = list(rag_chunk_pages(page_texts))

        elif ext in (".docx", ".doc"):
            text   = read_docx(tmp_path)
            chunks = list(rag_chunk_text(text))

        elif ext in (".xlsx", ".xls", ".xlsm", ".ods"):
            sheets = read_excel_sheets(tmp_path)
            chunks = list(rag_chunk_excel(sheets, source_name=original_name))

        elif ext in (".csv", ".tsv"):
            text   = read_csv(tmp_path, ext)
            chunks = list(rag_chunk_text(text))

        elif ext in (".jpg", ".jpeg", ".png", ".bmp", ".webp", ".tiff", ".tif"):
            from PIL import Image, ImageSequence
            img        = Image.open(tmp_path)
            frames     = list(ImageSequence.Iterator(img))
            page_texts = []
            for i, frame in enumerate(frames):
                b64 = to_base64(frame.convert("RGB"))
                r   = _client.chat(
                    model=MODEL_DEFAULT,
                    messages=[{"role": "user",
                               "content": f"Извлеки весь текст со страницы {i+1}. Верни только текст, без пояснений.",
                               "images": [b64]}],
                    options={"num_ctx": 4096, "temperature": 0.0},
                    stream=False,
                )
                page_texts.append(r["message"]["content"])
            if not page_texts:
                raise ValueError("Не удалось извлечь текст из изображения")
            chunks = list(rag_chunk_pages(page_texts))

        else:
            raise ValueError(f"Неподдерживаемый формат: {ext}")

        rag_index_document(doc_id, chunks)

    except Exception as e:
        with get_db() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "UPDATE kb_documents SET status='error', meta=%s WHERE id=%s",
                    (json.dumps({"error": str(e)}), doc_id)
                )
            conn.commit()
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


# ── статика ───────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return send_from_directory(ROOT_DIR, "index.html")


@app.route("/sso-login")
def sso_login():
    username = _normalize_username(request.args.get("u", ""))
    display_name = (request.args.get("d") or "").strip()
    ts_raw = (request.args.get("ts") or "").strip()
    sig = (request.args.get("sig") or "").strip()
    if not username or not display_name or not ts_raw or not sig:
        return "SSO parameters are missing", 400
    if not _verify_sso_payload(username, display_name, ts_raw, sig):
        return "SSO verification failed", 403
    session["logged_in"] = True
    session["username"] = username
    session["display_name"] = display_name
    return redirect("/")


@app.route("/api/me")
def api_me():
    username = _normalize_username(session.get("username") or "")
    display_name = (session.get("display_name") or username or "").strip()
    return jsonify({
        "logged_in": bool(session.get("logged_in")),
        "username": username,
        "display_name": display_name
    })


@app.route("/back.png")
def bg_image():
    return send_from_directory(ROOT_DIR, "back.png")

@app.route("/ico.png")
def ico_image():
    return send_from_directory(ROOT_DIR, "ico.png")

@app.route("/name.png")
def name_image():
    return send_from_directory(ROOT_DIR, "name.png")


@app.route("/api/models")
def api_models():
    available = get_available_models()
    result, seen_labels = [], set()
    for m in available:
        label = MODEL_LABELS.get(m, m)
        if label not in seen_labels:
            result.append({"id": m, "label": label})
            seen_labels.add(label)
    return {"models": result}


@app.route("/api/detect-file", methods=["POST"])
def api_detect_file():
    """Определяет тип файла и предлагает модель. Для PDF проверяет, скан это или текст."""
    file = request.files.get("file")
    if not file or not file.filename:
        return {"type": "unknown", "scanned": False, "model": MODEL_DEFAULT, "reason": ""}

    ext = os.path.splitext(file.filename)[1].lower()
    available = get_available_models()

    if ext != ".pdf":
        model, reason = suggest_model(ext, "", available)
        return {"type": ext.lstrip("."), "scanned": False, "model": model, "reason": reason}

    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp_path = tmp.name
    tmp.close()
    file.save(tmp_path)
    try:
        is_scanned = detect_pdf_type(tmp_path) == "scanned"
        eff_ext = ".jpg" if is_scanned else ".pdf"
        model, reason = suggest_model(eff_ext, "", available)
        if is_scanned:
            reason = "Сканированный PDF — нужна vision-модель"
        return {"type": "pdf", "scanned": is_scanned, "model": model, "reason": reason}
    finally:
        os.unlink(tmp_path)


@app.route("/api/suggest-model", methods=["POST"])
def api_suggest_model():
    data = request.get_json(force=True)
    file_ext = data.get("file_ext", "").lower()
    question = data.get("question", "")
    available = get_available_models()
    model, reason = suggest_model(file_ext, question, available)
    return {"model": model, "reason": reason}


# ── Сравнение документов: извлечение форматирования ─────────────────────────

def _to_cm(v):
    if v is None:
        return None
    try:
        return round(v.cm, 2)
    except AttributeError:
        try:
            return round(float(v) / 914400 * 2.54, 2)
        except Exception:
            return None


def _to_pt(v):
    if v is None:
        return None
    try:
        return round(v.pt, 1)
    except AttributeError:
        try:
            return round(float(v) / 12700, 1)
        except Exception:
            return None


def _build_page_numbers(doc) -> tuple:
    """
    Returns (page_list, has_info).
    page_list[i] = page number (1-based) for doc.paragraphs[i].
    has_info = True when the document contains reliable page-break markers.

    Detects three kinds of page breaks:
      1. paragraph_format.page_break_before  (always reliable)
      2. <w:br w:type="page"/> inside runs   (manual Ctrl+Enter breaks)
      3. <w:lastRenderedPageBreak/>           (soft breaks saved by Word)
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    paras = list(doc.paragraphs)
    n = len(paras)
    starts_new = [False] * n
    has_any = False

    # Check once whether the document has lastRenderedPageBreak info
    has_lrpb = any(
        run._r.find(f'{{{W}}}lastRenderedPageBreak') is not None
        for para in paras
        for run in para.runs
    )

    for i, para in enumerate(paras):
        # 1. page_break_before paragraph property
        try:
            if para.paragraph_format.page_break_before:
                starts_new[i] = True
                has_any = True
        except Exception:
            pass

        # 2. lastRenderedPageBreak — only counts when it appears BEFORE
        #    any real text in the paragraph (meaning the paragraph itself
        #    is on the new page, not just the tail of the previous one).
        if has_lrpb and not starts_new[i]:
            found_text = False
            done = False
            for run in para.runs:
                if done:
                    break
                for child in run._r:
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag == 'lastRenderedPageBreak':
                        if not found_text:
                            starts_new[i] = True
                            has_any = True
                        done = True
                        break
                    elif tag == 't' and (child.text or '').strip():
                        found_text = True

        # 3. Explicit manual page break inside THIS paragraph → NEXT para is on new page
        for run in para.runs:
            for br in run._r.findall(f'{{{W}}}br'):
                if br.get(f'{{{W}}}type') == 'page':
                    if i + 1 < n:
                        starts_new[i + 1] = True
                    has_any = True

    page, pages = 1, []
    for i, starts in enumerate(starts_new):
        if starts and i > 0:
            page += 1
        pages.append(page)

    return pages, (has_lrpb or has_any)


def _para_loc(p: dict) -> str:
    """Human-readable location string for a paragraph dict."""
    page = p.get('page')       # None when page info is unavailable
    num  = p.get('para_num', '?')
    return f"стр. {page}, абз. {num}" if page is not None else f"абз. {num}"


def extract_docx_formatting(path: str) -> dict:
    import docx as python_docx
    doc = python_docx.Document(path)
    ALIGN = {0: "слева", 1: "по центру", 2: "справа", 3: "по ширине"}

    margins = {}
    if doc.sections:
        s = doc.sections[0]
        margins = {
            "top":    _to_cm(s.top_margin),
            "bottom": _to_cm(s.bottom_margin),
            "left":   _to_cm(s.left_margin),
            "right":  _to_cm(s.right_margin),
        }

    doc_paras = list(doc.paragraphs)
    page_nums, has_page_info = _build_page_numbers(doc)

    paras = []
    para_num = 0

    for idx, para in enumerate(doc_paras):
        text = para.text.strip()
        if not text:
            continue
        para_num += 1
        pf = para.paragraph_format

        try:
            ls = pf.line_spacing
            ls_str = None
            if ls is not None:
                try:
                    ls_str = f"{ls.pt:.1f}pt"
                except AttributeError:
                    ls_str = f"{float(ls):.2f}x"
        except Exception:
            ls_str = None

        p_info = {
            "text":          text[:800],
            "page":          page_nums[idx] if has_page_info else None,
            "para_num":      para_num,
            "style":         para.style.name if para.style else None,
            "align":         ALIGN.get(para.alignment),
            "indent_left":   _to_cm(pf.left_indent),
            "indent_first":  _to_cm(pf.first_line_indent),
            "space_before":  _to_pt(pf.space_before),
            "space_after":   _to_pt(pf.space_after),
            "line_spacing":  ls_str,
            "runs": [],
        }

        for run in para.runs:
            rt = run.text.strip()
            if not rt:
                continue
            f = run.font
            size = _to_pt(f.size)
            attrs = []
            if run.bold:      attrs.append("жирный")
            if run.italic:    attrs.append("курсив")
            if run.underline: attrs.append("подч.")
            p_info["runs"].append({
                "text":    rt[:60],
                "font":    f.name,
                "size_pt": size,
                "attrs":   attrs,
            })

        paras.append(p_info)

    return {"margins": margins, "paragraphs": paras}


def _fmt_docx_for_prompt(name: str, data: dict, max_paras: int = 80) -> str:
    lines = [f"📄 {name}"]

    m = data.get("margins", {})
    if any(v is not None for v in m.values()):
        parts = []
        for k, lbl in [("top", "верх"), ("bottom", "низ"), ("left", "лево"), ("right", "право")]:
            if m.get(k) is not None:
                parts.append(f"{lbl}={m[k]}см")
        lines.append("Поля: " + ", ".join(parts))

    lines.append("")
    all_paras = data.get("paragraphs", [])
    paras = all_paras[:max_paras]

    for i, p in enumerate(paras, 1):
        # Полный текст абзаца
        full_text = p["text"].replace("\n", " ")
        lines.append(f"[{i}] {full_text}")

        # Параметры форматирования — одной строкой
        fmt = []
        if p.get("style"):                    fmt.append(f"стиль={p['style']}")
        if p.get("align"):                    fmt.append(f"выравн={p['align']}")
        if p.get("indent_left")  is not None: fmt.append(f"отст.лево={p['indent_left']}см")
        if p.get("indent_first") is not None: fmt.append(f"красн.стр={p['indent_first']}см")
        if p.get("space_before") is not None: fmt.append(f"до={p['space_before']}pt")
        if p.get("space_after")  is not None: fmt.append(f"после={p['space_after']}pt")
        if p.get("line_spacing"):             fmt.append(f"интервал={p['line_spacing']}")
        if fmt:
            lines.append("  [" + ", ".join(fmt) + "]")

        # Шрифты — уникальные комбинации
        seen_fonts = []
        for r in p.get("runs", []):
            parts = []
            if r.get("font"):    parts.append(r["font"])
            if r.get("size_pt"): parts.append(f"{r['size_pt']}pt")
            if r.get("attrs"):   parts.extend(r["attrs"])
            key = " ".join(parts)
            if key and key not in seen_fonts:
                seen_fonts.append(key)
        if seen_fonts:
            lines.append("  [шрифт: " + "; ".join(seen_fonts[:3]) + "]")

        lines.append("")

    if len(all_paras) > max_paras:
        lines.append(f"... ещё {len(all_paras) - max_paras} абзацев (не показаны)")

    return "\n".join(lines)


def extract_pdf_formatting(path: str) -> dict:
    """Extract per-block formatting from a text PDF via pymupdf dict mode."""
    import pymupdf
    from collections import Counter

    PT_TO_CM = 1 / 28.35
    doc = pymupdf.open(path)
    all_paragraphs = []
    page_margins = {}

    for page_idx, page in enumerate(doc):
        pw = page.rect.width
        ph = page.rect.height
        blocks = page.get_text("dict")["blocks"]
        text_blocks = [b for b in blocks if b.get("type") == 0]

        if text_blocks and page_idx == 0:
            xs0 = [b["bbox"][0] for b in text_blocks]
            xs1 = [b["bbox"][2] for b in text_blocks]
            ys0 = [b["bbox"][1] for b in text_blocks]
            ys1 = [b["bbox"][3] for b in text_blocks]
            page_margins = {
                "top":    round(min(ys0) * PT_TO_CM, 2),
                "bottom": round((ph - max(ys1)) * PT_TO_CM, 2),
                "left":   round(min(xs0) * PT_TO_CM, 2),
                "right":  round((pw - max(xs1)) * PT_TO_CM, 2),
            }

        for block in text_blocks:
            lines = block.get("lines", [])
            all_spans = [
                span
                for line in lines
                for span in line.get("spans", [])
                if span.get("text", "").strip()
            ]
            block_text = " ".join(s["text"] for s in all_spans).strip()
            if not block_text:
                continue

            fonts  = [s["font"] for s in all_spans]
            sizes  = [round(s["size"], 1) for s in all_spans]
            flags  = [s.get("flags", 0) for s in all_spans]

            dom_font  = Counter(fonts).most_common(1)[0][0] if fonts else None
            dom_size  = Counter(sizes).most_common(1)[0][0] if sizes else None
            is_bold   = any(f & 16 for f in flags)
            is_italic = any(f & 2  for f in flags)

            # Red-line indent: first line x0 vs average of rest
            indent_first = None
            if len(lines) >= 2:
                first_spans = lines[0].get("spans", [])
                rest_x0s = [
                    line["spans"][0]["bbox"][0]
                    for line in lines[1:]
                    if line.get("spans")
                ]
                if first_spans and rest_x0s:
                    x0_first = first_spans[0]["bbox"][0]
                    avg_rest = sum(rest_x0s) / len(rest_x0s)
                    indent_first = round((x0_first - avg_rest) * PT_TO_CM, 2)

            # Line spacing: avg distance between line tops
            line_spacing = None
            if len(lines) >= 2:
                y_tops = [ln["bbox"][1] for ln in lines]
                gaps = [y_tops[i + 1] - y_tops[i] for i in range(len(y_tops) - 1)]
                if gaps:
                    line_spacing = f"{round(sum(gaps) / len(gaps) * PT_TO_CM, 2)}см"

            attrs = []
            if is_bold:   attrs.append("жирный")
            if is_italic: attrs.append("курсив")

            all_paragraphs.append({
                "text":         block_text[:800],
                "page":         page_idx + 1,
                "para_num":     len(all_paragraphs) + 1,
                "font":         dom_font,
                "size_pt":      dom_size,
                "attrs":        attrs,
                "indent_first": indent_first,
                "line_spacing": line_spacing,
            })

    return {"margins": page_margins, "paragraphs": all_paragraphs}


def _extract_doc_for_compare(path: str, ext: str, name: str) -> str:
    try:
        if ext in (".docx", ".doc"):
            return _fmt_docx_for_prompt(name, extract_docx_formatting(path))
        elif ext == ".pdf":
            pages = read_pdf_text(path) if detect_pdf_type(path) == "text" else []
            if pages:
                # Нумеруем страницы, чтобы можно было указать место расхождения
                numbered = "\n\n".join(f"[Стр. {i+1}]\n{p}" for i, p in enumerate(pages))
                text = numbered[:16000]
            else:
                text = "[Сканированный PDF — текст не извлечён]"
            return f"📄 {name} [PDF — только текст, форматирование недоступно]\n\n{text}"
        elif ext in (".txt", ".md", ".log"):
            raw = read_txt(path)
            # Разбиваем на абзацы с нумерацией для точного указания места
            paras = [p.strip() for p in raw.split("\n\n") if p.strip()]
            numbered = "\n\n".join(f"[{i+1}] {p}" for i, p in enumerate(paras[:120]))
            return f"📄 {name} [текст]\n\n{numbered}"
        else:
            return f"📄 {name} [формат {ext} не поддерживается]"
    except Exception as e:
        return f"📄 {name} [ошибка извлечения: {e}]"


# ── Сравнение документов: алгоритмический diff ───────────────────────────────

def _word_diff(orig_text: str, copy_text: str) -> list:
    a, b = orig_text.split(), copy_text.split()
    if not a and not b:
        return []
    sm = difflib.SequenceMatcher(None, a, b, autojunk=False)
    changes = []
    for tag, a1, a2, b1, b2 in sm.get_opcodes():
        if tag == 'replace':
            changes.append(f'«{" ".join(a[a1:a2])}» → «{" ".join(b[b1:b2])}»')
        elif tag == 'delete':
            changes.append(f'удалено «{" ".join(a[a1:a2])}»')
        elif tag == 'insert':
            changes.append(f'добавлено «{" ".join(b[b1:b2])}»')
    return changes


def _dominant_font(runs: list) -> str:
    counts = {}
    for r in runs:
        parts = []
        if r.get('font'):    parts.append(r['font'])
        if r.get('size_pt'): parts.append(f"{r['size_pt']}pt")
        if r.get('attrs'):   parts.extend(r['attrs'])
        key = " ".join(parts) or '—'
        counts[key] = counts.get(key, 0) + len(r.get('text', ''))
    return max(counts, key=counts.get) if counts else '—'


def _para_fmt_diff(orig: dict, copy: dict) -> list:
    result = []
    loc = _para_loc(orig)
    for key, lbl in [('style', 'Стиль'), ('align', 'Выравнивание'),
                     ('indent_left', 'Отступ слева (см)'),
                     ('indent_first', 'Красная строка (см)'),
                     ('space_before', 'Интервал до (pt)'),
                     ('space_after', 'Интервал после (pt)'),
                     ('line_spacing', 'Межстрочный интервал')]:
        ov, cv = orig.get(key), copy.get(key)
        if ov != cv:
            result.append({
                'loc': loc,
                'param': lbl,
                'orig': str(ov) if ov is not None else '(не задано)',
                'copy': str(cv) if cv is not None else '(не задано)',
                'text': orig['text'][:70],
            })
    of = _dominant_font(orig.get('runs', []))
    cf = _dominant_font(copy.get('runs', []))
    if of != cf:
        result.append({'loc': loc, 'param': 'Шрифт',
                       'orig': of, 'copy': cf, 'text': orig['text'][:70]})
    return result


def _compare_docx_data(orig_data: dict, copy_data: dict) -> dict:
    margin_diffs, text_diffs, fmt_diffs = [], [], []

    om = orig_data.get('margins', {})
    cm_m = copy_data.get('margins', {})
    for key, lbl in [('top', 'Верхнее поле'), ('bottom', 'Нижнее поле'),
                     ('left', 'Левое поле'), ('right', 'Правое поле')]:
        ov, cv = om.get(key), cm_m.get(key)
        if ov != cv:
            margin_diffs.append({
                'param': lbl,
                'orig': f'{ov} см' if ov is not None else 'н/д',
                'copy': f'{cv} см' if cv is not None else 'н/д',
            })

    op = orig_data.get('paragraphs', [])
    cp = copy_data.get('paragraphs', [])
    sm = difflib.SequenceMatcher(None,
                                  [p['text'] for p in op],
                                  [p['text'] for p in cp],
                                  autojunk=False)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            for oi, ci in zip(range(i1, i2), range(j1, j2)):
                fmt_diffs.extend(_para_fmt_diff(op[oi], cp[ci]))
        elif tag == 'replace':
            if i2 - i1 == 1 and j2 - j1 == 1:
                changes = _word_diff(op[i1]['text'], cp[j1]['text'])
                if changes:
                    text_diffs.append({
                        'type': 'changed',
                        'loc': _para_loc(op[i1]),
                        'changes': changes[:15],
                        'orig': op[i1]['text'][:150],
                        'copy': cp[j1]['text'][:150],
                    })
                fmt_diffs.extend(_para_fmt_diff(op[i1], cp[j1]))
            else:
                for pi in range(i1, i2):
                    text_diffs.append({'type': 'deleted', 'loc': _para_loc(op[pi]),
                                       'text': op[pi]['text'][:150]})
                for pj in range(j1, j2):
                    text_diffs.append({'type': 'added', 'loc': _para_loc(cp[pj]),
                                       'text': cp[pj]['text'][:150]})
        elif tag == 'delete':
            for pi in range(i1, i2):
                text_diffs.append({'type': 'deleted', 'loc': _para_loc(op[pi]),
                                   'text': op[pi]['text'][:150]})
        elif tag == 'insert':
            for pj in range(j1, j2):
                text_diffs.append({'type': 'added', 'loc': _para_loc(cp[pj]),
                                   'text': cp[pj]['text'][:150]})

    return {'margin_diffs': margin_diffs, 'text_diffs': text_diffs, 'fmt_diffs': fmt_diffs}


def _compare_pdf_data(orig_data: dict, copy_data: dict) -> dict:
    """Full diff for two text PDFs: margins + text + per-block formatting."""
    margin_diffs, text_diffs, fmt_diffs = [], [], []

    om   = orig_data.get("margins", {})
    cm_m = copy_data.get("margins", {})
    for key, lbl in [("top", "Верхнее поле"), ("bottom", "Нижнее поле"),
                     ("left", "Левое поле"), ("right", "Правое поле")]:
        ov, cv = om.get(key), cm_m.get(key)
        if ov != cv:
            margin_diffs.append({
                "param": lbl,
                "orig":  f"{ov} см" if ov is not None else "н/д",
                "copy":  f"{cv} см" if cv is not None else "н/д",
            })

    op = orig_data.get("paragraphs", [])
    cp = copy_data.get("paragraphs", [])
    sm = difflib.SequenceMatcher(None,
                                  [p["text"] for p in op],
                                  [p["text"] for p in cp],
                                  autojunk=False)

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for oi, ci in zip(range(i1, i2), range(j1, j2)):
                loc = _para_loc(op[oi])
                for key, lbl in [("font", "Шрифт"), ("size_pt", "Размер (pt)"),
                                  ("indent_first", "Красная строка (см)"),
                                  ("line_spacing", "Межстрочный интервал")]:
                    ov = op[oi].get(key)
                    cv = cp[ci].get(key)
                    if ov != cv:
                        fmt_diffs.append({
                            "loc":   loc,
                            "param": lbl,
                            "orig":  str(ov) if ov is not None else "(не задано)",
                            "copy":  str(cv) if cv is not None else "(не задано)",
                            "text":  op[oi]["text"][:70],
                        })
                # attrs is a list — compare as sorted tuple
                oa = ", ".join(sorted(op[oi].get("attrs", []))) or "обычный"
                ca = ", ".join(sorted(cp[ci].get("attrs", []))) or "обычный"
                if oa != ca:
                    fmt_diffs.append({
                        "loc": loc, "param": "Начертание",
                        "orig": oa, "copy": ca, "text": op[oi]["text"][:70],
                    })
        elif tag == "replace":
            if i2 - i1 == 1 and j2 - j1 == 1:
                changes = _word_diff(op[i1]["text"], cp[j1]["text"])
                if changes:
                    text_diffs.append({
                        "type": "changed", "loc": _para_loc(op[i1]),
                        "changes": changes[:15],
                        "orig": op[i1]["text"][:150], "copy": cp[j1]["text"][:150],
                    })
            else:
                for pi in range(i1, i2):
                    text_diffs.append({"type": "deleted", "loc": _para_loc(op[pi]),
                                       "text": op[pi]["text"][:150]})
                for pj in range(j1, j2):
                    text_diffs.append({"type": "added", "loc": _para_loc(cp[pj]),
                                       "text": cp[pj]["text"][:150]})
        elif tag == "delete":
            for pi in range(i1, i2):
                text_diffs.append({"type": "deleted", "loc": _para_loc(op[pi]),
                                   "text": op[pi]["text"][:150]})
        elif tag == "insert":
            for pj in range(j1, j2):
                text_diffs.append({"type": "added", "loc": _para_loc(cp[pj]),
                                   "text": cp[pj]["text"][:150]})

    return {"margin_diffs": margin_diffs, "text_diffs": text_diffs, "fmt_diffs": fmt_diffs}


def _compare_text_data(orig_text: str, copy_text: str) -> dict:
    def split_paras(text):
        """Split into dicts with text/page/para_num; honours [Страница N] markers from PDF."""
        current_page = 1
        para_num = 0
        result = []
        for chunk in re.split(r'\n{2,}', text):
            chunk = chunk.strip()
            if not chunk:
                continue
            m = re.match(r'^\[Страница (\d+)\]\n?(.*)', chunk, re.DOTALL)
            if m:
                current_page = int(m.group(1))
                body = m.group(2).strip()
                if body:
                    para_num += 1
                    result.append({'text': body, 'page': current_page, 'para_num': para_num})
            else:
                para_num += 1
                result.append({'text': chunk, 'page': current_page, 'para_num': para_num})
        return result

    op = split_paras(orig_text)
    cp = split_paras(copy_text)
    text_diffs = []
    sm = difflib.SequenceMatcher(None,
                                  [p['text'] for p in op],
                                  [p['text'] for p in cp],
                                  autojunk=False)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'replace':
            if i2 - i1 == 1 and j2 - j1 == 1:
                changes = _word_diff(op[i1]['text'], cp[j1]['text'])
                if changes:
                    text_diffs.append({
                        'type': 'changed', 'loc': _para_loc(op[i1]),
                        'changes': changes[:15],
                        'orig': op[i1]['text'][:150], 'copy': cp[j1]['text'][:150],
                    })
            else:
                for pi in range(i1, i2):
                    text_diffs.append({'type': 'deleted', 'loc': _para_loc(op[pi]),
                                       'text': op[pi]['text'][:150]})
                for pj in range(j1, j2):
                    text_diffs.append({'type': 'added', 'loc': _para_loc(cp[pj]),
                                       'text': cp[pj]['text'][:150]})
        elif tag == 'delete':
            for pi in range(i1, i2):
                text_diffs.append({'type': 'deleted', 'loc': _para_loc(op[pi]),
                                   'text': op[pi]['text'][:150]})
        elif tag == 'insert':
            for pj in range(j1, j2):
                text_diffs.append({'type': 'added', 'loc': _para_loc(cp[pj]),
                                   'text': cp[pj]['text'][:150]})
    return {'margin_diffs': [], 'text_diffs': text_diffs, 'fmt_diffs': []}


def _render_diff_report(copy_name: str, diff: dict) -> str:
    md_d = diff['margin_diffs']
    td   = diff['text_diffs']
    fd   = diff['fmt_diffs']
    total = len(md_d) + len(td) + len(fd)

    lines = [f"## Копия — «{copy_name}»", ""]
    if total == 0:
        lines.append("✅ **Документ полностью соответствует оригиналу — расхождений не обнаружено.**")
        return '\n'.join(lines)

    lines.append(
        f"Обнаружено расхождений: **{total}** "
        f"(текстовых: **{len(td)}**, "
        f"форматирование: **{len(md_d) + len(fd)}**)"
    )
    lines.append("")

    if md_d:
        lines += ["### 📐 Поля страницы", ""]
        for d in md_d:
            lines.append(f"- **{d['param']}:** `{d['orig']}` → `{d['copy']}`")
        lines.append("")

    if td:
        lines += ["### ✏️ Текстовые расхождения", ""]
        for d in td:
            loc = d.get('loc', '?')
            if d['type'] == 'deleted':
                lines.append(f"- ❌ **Удалён** ({loc}): «{d['text']}»")
            elif d['type'] == 'added':
                lines.append(f"- ➕ **Добавлен** ({loc}): «{d['text']}»")
            elif d['type'] == 'changed':
                lines.append(f"- 🔄 **Изменён** ({loc}):")
                lines.append(f"  - Оригинал: «{d['orig']}»")
                lines.append(f"  - Копия:    «{d['copy']}»")
                lines.append(f"  - Изменения:")
                for ch in d['changes']:
                    lines.append(f"    - {ch}")
        lines.append("")

    if fd:
        lines += ["### 🎨 Форматирование абзацев", ""]
        for d in fd:
            lines.append(
                f"- **{d['param']}** ({d.get('loc', '?')}, «{d['text']}…»): "
                f"`{d['orig']}` → `{d['copy']}`"
            )
        lines.append("")

    return '\n'.join(lines)


@app.route("/api/compare", methods=["POST"])
def api_compare_documents():
    model   = request.form.get("model", MODEL_DEFAULT)
    orig_f  = request.files.get("original")
    copy_fs = request.files.getlist("copies")

    if not orig_f or not orig_f.filename:
        return jsonify({"error": "Оригинальный документ не передан"}), 400
    if not copy_fs:
        return jsonify({"error": "Не переданы копии для сравнения"}), 400

    def _save(f):
        ext = os.path.splitext(f.filename)[1].lower()
        tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
        p = tmp.name
        tmp.close()
        f.save(p)
        return p, ext

    orig_path, orig_ext = _save(orig_f)
    copies_saved = [_save(f) for f in copy_fs]
    copy_names = [f.filename for f in copy_fs]

    def generate():
        try:
            # Pre-extract original data once
            profile = get_profile(model)
            if orig_ext in ('.docx', '.doc'):
                orig_docx_data = extract_docx_formatting(orig_path)
                orig_pdf_data  = None
                orig_plain_text = "\n\n".join(
                    p['text'] for p in orig_docx_data.get('paragraphs', [])
                )
            elif orig_ext == '.pdf':
                orig_docx_data = None
                if detect_pdf_type(orig_path) == "text":
                    orig_pdf_data   = extract_pdf_formatting(orig_path)
                    orig_plain_text = "\n\n".join(
                        p['text'] for p in orig_pdf_data.get('paragraphs', [])
                    )
                else:
                    orig_pdf_data = None
                    yield f"data: {json.dumps({'text': '🖼️ Оригинал — скан, распознаю текст...\\n\\n'})}\n\n"
                    orig_plain_text = read_scanned_pdf_as_text(orig_path, model, profile)
            else:
                orig_plain_text = read_txt(orig_path)
                orig_docx_data  = None
                orig_pdf_data   = None

            for i, (copy_path, copy_ext) in enumerate(copies_saved):
                copy_name = copy_names[i]
                if i > 0:
                    yield f"data: {json.dumps({'text': chr(10) + '---' + chr(10) + chr(10)})}\n\n"

                try:
                    if orig_docx_data is not None and copy_ext in ('.docx', '.doc'):
                        copy_docx_data = extract_docx_formatting(copy_path)
                        diff = _compare_docx_data(orig_docx_data, copy_docx_data)
                    elif orig_pdf_data is not None and copy_ext == '.pdf':
                        if detect_pdf_type(copy_path) == "text":
                            copy_pdf_data = extract_pdf_formatting(copy_path)
                            diff = _compare_pdf_data(orig_pdf_data, copy_pdf_data)
                        else:
                            yield f"data: {json.dumps({'text': f'🖼️ Копия — скан, распознаю текст...\\n\\n'})}\n\n"
                            copy_plain = read_scanned_pdf_as_text(copy_path, model, profile)
                            diff = _compare_text_data(orig_plain_text, copy_plain)
                    else:
                        if copy_ext in ('.docx', '.doc'):
                            copy_plain = read_docx(copy_path)
                        elif copy_ext == '.pdf':
                            if detect_pdf_type(copy_path) == "text":
                                pages = read_pdf_text(copy_path)
                                copy_plain = "\n\n".join(pages)
                            else:
                                yield f"data: {json.dumps({'text': f'🖼️ Копия — скан, распознаю текст...\\n\\n'})}\n\n"
                                copy_plain = read_scanned_pdf_as_text(copy_path, model, profile)
                        else:
                            copy_plain = read_txt(copy_path)
                        diff = _compare_text_data(orig_plain_text, copy_plain)

                    report = _render_diff_report(copy_name, diff)
                    for line in report.split('\n'):
                        yield f"data: {json.dumps({'text': line + chr(10)})}\n\n"

                except Exception as e:
                    yield f"data: {json.dumps({'text': chr(10) + '⚠️ Ошибка при сравнении «' + copy_name + '»: ' + str(e) + chr(10)})}\n\n"

            yield "data: [DONE]\n\n"

        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
            yield "data: [DONE]\n\n"
        finally:
            for p in [orig_path] + [p for p, _ in copies_saved]:
                try:
                    os.unlink(p)
                except OSError:
                    pass

    return Response(stream_with_context(generate()), mimetype="text/event-stream")


if __name__ == "__main__":
    # init_db()
    print("🚀 Запуск на http://localhost:5000")
    app.run(host="0.0.0.0", port=5000, debug=False, threaded=True)
