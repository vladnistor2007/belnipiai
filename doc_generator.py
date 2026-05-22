import os
import re
import sys
import uuid
import subprocess
import tempfile

TEMP_DIR = os.path.join(tempfile.gettempdir(), "belnipiai_docs")
EXEC_TIMEOUT = 30  # seconds

_FONT_CANDIDATES = [
    # Linux
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    "/usr/share/fonts/liberation/LiberationSans-Regular.ttf",
    "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
    "/usr/share/fonts/truetype/ubuntu/Ubuntu-R.ttf",
    # Windows
    r"C:\Windows\Fonts\arial.ttf",
    r"C:\Windows\Fonts\calibri.ttf",
    r"C:\Windows\Fonts\times.ttf",
]


def _find_cyrillic_font() -> str:
    for path in _FONT_CANDIDATES:
        if os.path.exists(path):
            return path
    return ""

_ACTION_WORDS = [
    "создай", "сформируй", "напиши", "сгенерируй",
    "составь", "подготовь", "сделай", "создать",
]
_FORMAT_WORDS = [
    "word", "ворд", "docx",
    "excel", "эксель", "xlsx",
    "pdf", "пдф",
    "документ", "файл",
]

SYSTEM_PROMPT_DOCGEN = """\
Ты — ассистент по созданию файлов (Word, Excel, PDF) на предприятии.
Твоя единственная задача — написать Python-код, который создаёт файл.

АБСОЛЮТНЫЕ ПРАВИЛА — нарушать нельзя:
1. ВСЕГДА пиши Python-код в блоке ```python:document ... ```
2. НИКОГДА не показывай данные в виде markdown-таблиц, списков или текста — только код
3. НИКОГДА не пиши "скопируй и вставь", "вот данные", "содержимое файла" — только код
4. Переменная OUTPUT_PATH уже определена — файл сохраняй только в неё
5. Код должен быть полным и рабочим — все данные вшиты прямо в код, без заглушек
6. До блока кода — ровно одно предложение что создаёшь, больше ничего
7. Доступные библиотеки:
   Word  → from docx import Document; from docx.shared import Pt, Inches, RGBColor; from docx.enum.text import WD_ALIGN_PARAGRAPH
   Excel → import openpyxl; from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
   PDF   → from fpdf import FPDF
   Прочие → os, datetime, re, json, math, collections
8. Не вызывай pip, subprocess, не читай произвольные файлы системы

ВАЖНО ДЛЯ EXCEL: данные из запроса пользователя вписывай прямо в код как Python-списки или словари. Пример: данные = [["Год", "Знак"], [1925, "Овен"], ...]; для строки в данных: ws.append(строка)

ШАБЛОН WORD:
```python:document
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

h = doc.add_heading('', 0)
run = h.add_run('Заголовок документа')
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph('Текст параграфа.')

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
for i, name in enumerate(['Колонка 1', 'Колонка 2', 'Колонка 3']):
    table.rows[0].cells[i].text = name
row = table.add_row().cells
row[0].text = 'Значение'

doc.save(OUTPUT_PATH)
```

ШАБЛОН EXCEL:
```python:document
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill

wb = openpyxl.Workbook()
ws = wb.active
ws.title = 'Данные'

ws['A1'] = 'Заголовок'
ws['A1'].font = Font(bold=True, size=12, color='1F497D')
ws['A1'].fill = PatternFill('solid', fgColor='DCE6F1')
ws.column_dimensions['A'].width = 25

wb.save(OUTPUT_PATH)
```

ШАБЛОН PDF (кириллица — FONT_PATH уже определён системой):
```python:document
from fpdf import FPDF

pdf = FPDF()
pdf.add_page()
pdf.set_auto_page_break(auto=True, margin=15)

if FONT_PATH:
    pdf.add_font('F', fname=FONT_PATH)
    pdf.set_font('F', size=14)
else:
    pdf.set_font('Helvetica', size=14)

pdf.cell(0, 12, 'Заголовок документа', new_x='LMARGIN', new_y='NEXT')
pdf.set_font_size(11)
pdf.multi_cell(0, 8, 'Текст документа.')
pdf.output(OUTPUT_PATH)
```
"""


def is_document_request(text: str) -> bool:
    t = text.lower()
    return any(a in t for a in _ACTION_WORDS) and any(f in t for f in _FORMAT_WORDS)


def detect_format(text: str) -> str:
    t = text.lower()
    if any(w in t for w in ("excel", "эксель", "xlsx")):
        return "xlsx"
    if any(w in t for w in ("pdf", "пдф")):
        return "pdf"
    return "docx"


def make_output_path(fmt: str) -> tuple[str, str]:
    os.makedirs(TEMP_DIR, exist_ok=True)
    name = f"{uuid.uuid4()}.{fmt}"
    return os.path.join(TEMP_DIR, name), name


def extract_code(text: str) -> str | None:
    for pattern in (
        r"```python:document\s*\n(.*?)```",
        r"```python\s*\n(.*?)```",
    ):
        m = re.search(pattern, text, re.DOTALL)
        if m:
            return m.group(1).strip()
    return None


def run_code(code: str, output_path: str) -> tuple[bool, str]:
    """Execute AI-generated document code in an isolated subprocess."""
    font_path = _find_cyrillic_font()
    full_code = (
        f"OUTPUT_PATH = {repr(output_path)}\n"
        f"FONT_PATH = {repr(font_path)}\n\n"
        f"{code}"
    )
    script = output_path + ".py"
    try:
        with open(script, "w", encoding="utf-8") as f:
            f.write(full_code)
        result = subprocess.run(
            [sys.executable, script],
            capture_output=True,
            text=True,
            timeout=EXEC_TIMEOUT,
            cwd=TEMP_DIR,
        )
        if os.path.exists(script):
            os.unlink(script)
        if result.returncode != 0:
            return False, (result.stderr or result.stdout or "Ошибка выполнения").strip()
        if not os.path.exists(output_path):
            return False, "Код выполнился, но файл не создан"
        return True, ""
    except subprocess.TimeoutExpired:
        if os.path.exists(script):
            os.unlink(script)
        return False, f"Превышено время ожидания ({EXEC_TIMEOUT} с)"
    except Exception as e:
        return False, str(e)
